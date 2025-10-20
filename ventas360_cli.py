#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Ventas 360 — CLI (Pandas)

Descripción
-----------
Analiza el dataset de ventas y genera salidas en la carpeta `out/`:
- CSV agregados por ciudad, producto, estado y mes.
- Resumen de KPIs (texto) y un informe analítico en Word (opcional).

Entradas
--------
- Archivo CSV principal en `data/ventas_5k.csv` (fechas 2021–2025).
- Montos en COP con 2 decimales.

Salidas (por defecto en `out/`)
-------------------------------
- CSV agregados (ej. `ventas_mensuales_por_ciudad.csv`, `tops_productos_por_registros.csv`, etc.).
- `resumen.txt` con KPIs clave.
- `Informe_Analitico_Ventas_Pandas.docx` (si se habilita).

Uso rápido
----------
python ventas360_cli.py --in data/ventas_5k.csv --out out --docx true

Parámetros
----------
--in     Ruta del CSV de entrada (default: data/ventas_5k.csv)
--out    Carpeta de salidas (default: out)
--docx   Generar informe Word (true/false, default: true)

Requisitos
----------
- Python 3.10+
- pip install -r requirements.txt

Notas
-----
- Las rutas se crean automáticamente si no existen.
- Los valores monetarios se formatean como COP con 2 decimales.
- Este script está pensado para correr antes del dashboard, de modo que el
  directorio `out/` quede poblado para su visualización en Streamlit.
"""


from __future__ import annotations

import argparse
from dataclasses import dataclass
from datetime import datetime, date
from pathlib import Path
from typing import Dict, Any, Optional, Tuple

import numpy as np
import pandas as pd

# -------------------- Config --------------------
RANDOM_SEED = 42

CATEGORIAS = [
    "Tecnología", "Hogar", "Alimentos", "Deportes", "Ropa",
    "Belleza", "Jardín", "Mascotas", "Salud", "Juguetes",
]

PRODUCTOS_POR_CATEGORIA = {
    "Tecnología": ["Laptop", "Tablet", "Celular", "Monitor", "Teclado", "Mouse", "Televisor"],
    "Hogar": ["Sofá", "Mesa", "Silla", "Cortinas", "Lámpara"],
    "Alimentos": ["Arroz", "Atún", "Aceite", "Leche", "Galletas"],
    "Deportes": ["Balón", "Bicicleta", "Pesas", "Proteína"],
    "Ropa": ["Pantalón", "Camisa", "Zapatos", "Chaqueta"],
    "Belleza": ["Perfume", "Maquillaje", "Crema"],
    "Jardín": ["Pala", "Manguera", "Abono"],
    "Mascotas": ["Concentrado", "Arena", "Juguete"],
    "Salud": ["Alcohol", "Algodón", "Vitaminas"],
    "Juguetes": ["Rompecabezas", "Muñeco", "Carritos"],
}

CIUDADES = [
    "Bogotá", "Medellín", "Cali", "Barranquilla", "Bucaramanga",
    "Cartagena", "Pereira", "Manizales", "Cúcuta", "Santa Marta",
]

VENDEDORES = [f"Vendedor_{i:03d}" for i in range(1, 201)]  # 200 vendedores

# rangos de precio por categoría (min, max)
RANGO_PRECIO = {
    "Tecnología": (300_000, 5_000_000),
    "Hogar": (50_000, 3_000_000),
    "Alimentos": (5_000, 200_000),
    "Deportes": (30_000, 1_500_000),
    "Ropa": (20_000, 800_000),
    "Belleza": (15_000, 600_000),
    "Jardín": (10_000, 700_000),
    "Mascotas": (10_000, 600_000),
    "Salud": (10_000, 600_000),
    "Juguetes": (10_000, 500_000),
}

# -------------------- Generación de datos --------------------
def fmt_cop(x) -> str:
    """Formatea número a COP: $ 1.234.567,89"""
    try:
        s = f"{float(x):,.2f}"
        s = s.replace(",", "_").replace(".", ",").replace("_", ".")
        return f"$ {s}"
    except Exception:
        return str(x)


def generar_dataset(n: int, csv_path: Path, seed: int = RANDOM_SEED) -> Path:
    """Genera un dataset sintético con n registros y lo guarda en csv_path."""
    rng = np.random.default_rng(seed)

    categorias = rng.choice(CATEGORIAS, size=n)
    productos = np.array([rng.choice(PRODUCTOS_POR_CATEGORIA[cat]) for cat in categorias], dtype=object)

    precios = np.empty(n, dtype=float)
    for cat in set(CATEGORIAS):
        low, high = RANGO_PRECIO[cat]
        mask = categorias == cat
        precios[mask] = rng.uniform(low, high, size=mask.sum())

    cantidades = rng.integers(1, 6, size=n)  # 1 a 5
    ciudades = rng.choice(CIUDADES, size=n)
    vendedores = rng.choice(VENDEDORES, size=n)

    # rango de fechas: del 2021-01-01 hasta HOY (sin futuros)
    start = np.datetime64("2021-01-01T00:00:00")
    today = date.today().isoformat()
    end = np.datetime64(f"{today}T23:59:00")  # hoy a las 23:59
    # tamaño del rango en MINUTOS (evita desbordes)
    delta_min = int(((end - start) / np.timedelta64(1, "m")))
    offset_min = rng.integers(0, delta_min + 1, size=n)
    fechas = start + offset_min.astype("timedelta64[m]")

    estados = rng.choice(["Cerrado", "Abierto", "Cancelado", "En Proceso"], size=n, p=[0.6, 0.2, 0.1, 0.1])
    comisiones_pct = rng.uniform(0.03, 0.08, size=n)

    valor_venta = (precios * cantidades).round(2)
    comision = (valor_venta * comisiones_pct).round(2)

    clientes = np.array([f"Cliente_{i:04d}" for i in rng.integers(1, 5000, size=n)], dtype=object)

    df = pd.DataFrame({
        "CLIENTE": clientes,
        "PRODUCTO": productos,
        "PRECIO": precios.round(2),
        "CANTIDAD": cantidades,
        "CIUDAD": ciudades,
        "VENDEDOR": vendedores,
        "FECHA": fechas.astype("datetime64[m]").astype(str),
        "ESTADO": estados,
        "VALOR_VENTA": valor_venta,
        "COMISION": comision,
        "CATEGORIA": categorias,
    })

    csv_path = Path(csv_path)
    csv_path.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(csv_path, index=False)
    return csv_path

# -------------------- Carga / normalización --------------------
def cargar_csv(csv_path: Path) -> pd.DataFrame:
    df = pd.read_csv(csv_path)
    # normalizar cabeceras
    mapping = {
        "VALOR_VENTA": "valor_venta",
        "UTILIDAD": "utilidad",
        "CATEGORIA": "categoria",
        "CIUDAD": "ciudad",
        "VENDEDOR": "vendedor",
        "ESTADO": "estado",
        "PRODUCTO": "producto",
        "FECHA": "fecha",
        "CLIENTE": "cliente",
        "PRECIO": "precio",
        "CANTIDAD": "cantidad",
        "COMISION": "comision",
    }
    df = df.rename(columns={k: v for k, v in mapping.items() if k in df.columns})
    # fecha robusta
    if "fecha" in df.columns:
        df["fecha"] = pd.to_datetime(df["fecha"], errors="coerce")
    return df

# -------------------- Análisis principal (enunciado 1–18) --------------------
def responder_preguntas(df: pd.DataFrame) -> Dict[str, Any]:
    """Resuelve las preguntas 1–18 del enunciado y retorna un dict de resultados."""
    resultados: Dict[str, Any] = {}
    d = df.copy()

    # 1 Total de registros
    resultados["total_registros"] = int(len(d))

    # 2 Ventas por estado
    if "estado" in d.columns:
        resultados["ventas_por_estado"] = d["estado"].value_counts()
    else:
        resultados["ventas_por_estado"] = pd.Series(dtype=int)

    # 3 Valor total de ventas
    resultados["valor_total_ventas"] = float(d.get("valor_venta", pd.Series(dtype=float)).sum())

    # 4 Comisión promedio en cerradas
    cerr = d[d.get("estado", pd.Series(index=d.index)) == "Cerrado"]
    if not cerr.empty and "comision" in cerr.columns:
        resultados["prom_comision_cerradas"] = float(cerr["comision"].mean())
    else:
        resultados["prom_comision_cerradas"] = 0.0

    # 5 Ciudad con más cerradas
    if not cerr.empty and "ciudad" in cerr.columns:
        resultados["ciudad_mas_cerradas"] = cerr["ciudad"].value_counts().idxmax()
    else:
        resultados["ciudad_mas_cerradas"] = None

    # 6 Valor total por ciudad
    if "ciudad" in d.columns and "valor_venta" in d.columns:
        resultados["ventas_por_ciudad_sum"] = (
            d.groupby("ciudad", as_index=True)["valor_venta"].sum().sort_values(ascending=False)
        )
    else:
        resultados["ventas_por_ciudad_sum"] = pd.Series(dtype=float)

    # 7 Top 5 productos por registros
    if "producto" in d.columns:
        resultados["top5_productos_por_registros"] = d["producto"].value_counts().head(5)
    else:
        resultados["top5_productos_por_registros"] = pd.Series(dtype=int)

    # 8 Productos únicos
    resultados["productos_unicos"] = int(d.get("producto", pd.Series(dtype=object)).nunique())

    # 9 Vendedor con más cerradas
    if not cerr.empty and "vendedor" in cerr.columns:
        resultados["vendedor_mas_cerradas"] = cerr["vendedor"].value_counts().idxmax()
    else:
        resultados["vendedor_mas_cerradas"] = None

    # 10 Venta de mayor valor (fila)
    if "valor_venta" in d.columns:
        idx_max = d["valor_venta"].idxmax()
        cols = [c for c in ["cliente", "producto", "valor_venta", "ciudad", "vendedor", "fecha"] if c in d.columns]
        resultados["venta_max"] = d.loc[idx_max, cols]
    else:
        resultados["venta_max"] = pd.Series(dtype=object)

    # 11 Media del VALOR_VENTA por mes
    if "fecha" in d.columns and "valor_venta" in d.columns and pd.api.types.is_datetime64_any_dtype(d["fecha"]):
        d["mes"] = d["fecha"].dt.to_period("M").dt.to_timestamp()
        resultados["media_valor_venta_por_mes"] = d.groupby("mes")["valor_venta"].mean().round(2)
    else:
        resultados["media_valor_venta_por_mes"] = pd.Series(dtype=float)

    # 12 Ventas por trimestre (conteo)
    if "fecha" in d.columns and pd.api.types.is_datetime64_any_dtype(d["fecha"]):
        resultados["ventas_por_trimestre"] = d["fecha"].dt.to_period("Q").value_counts().sort_index().to_dict()
    else:
        resultados["ventas_por_trimestre"] = {}

    # 13 Chequeos de calidad (valores negativos / nulos)
    v = d.get("valor_venta", pd.Series(dtype=float))
    c = d.get("comision", pd.Series(dtype=float))
    resultados["ventas_valor_negativo"] = int((v <= 0).sum()) if len(v) else 0
    resultados["comision_negativa"] = int((c < 0).sum()) if len(c) else 0
    resultados["nulos_valor"] = int(v.isna().sum()) if len(v) else 0
    resultados["nulos_comision"] = int(c.isna().sum()) if len(c) else 0

    # 14 Ventas mensuales por ciudad (suma)
    if "fecha" in d.columns and "ciudad" in d.columns and "valor_venta" in d.columns and pd.api.types.is_datetime64_any_dtype(d["fecha"]):
        d["Mes"] = d["fecha"].dt.to_period("M").astype(str)
        resultados["ventas_mensuales_por_ciudad"] = (
            d.groupby(["ciudad", "Mes"])["valor_venta"].sum().reset_index().rename(columns={"ciudad": "CIUDAD", "valor_venta": "VALOR_VENTA"})
        )
    else:
        resultados["ventas_mensuales_por_ciudad"] = pd.DataFrame(columns=["CIUDAD", "Mes", "VALOR_VENTA"])

    # 15 productos vendidos en más de 3 ciudades
    if "producto" in d.columns and "ciudad" in d.columns:
        prod_3_ciudades = d.groupby("producto")["ciudad"].nunique()
        resultados["productos_mas_de_3_ciudades"] = prod_3_ciudades[prod_3_ciudades > 3].sort_values(ascending=False)
    else:
        resultados["productos_mas_de_3_ciudades"] = pd.Series(dtype=int)

    # 16 duplicados
    resultados["duplicados_total"] = int(d.duplicated().sum())
    resultados["duplicados_ejemplo"] = d[d.duplicated()].head(10)

    # 17 limpieza por nulos (ejemplo simple)
    base_len = len(d)
    cols_required = [c for c in ["cliente", "producto", "valor_venta"] if c in d.columns]
    d_limpio = d.dropna(subset=cols_required) if cols_required else d
    resultados["filas_eliminadas_por_nulos"] = int(base_len - len(d_limpio))

    # 18 utilidad (enunciado: 95% del valor_venta) + KPIs aclaratorios
    if "valor_venta" in d.columns:
        d["UTILIDAD"] = (d["valor_venta"] * 0.95).round(2)  # enunciado
        d["COSTO"] = (d["valor_venta"] * 0.95).round(2)
        d["UTILIDAD_5PCT"] = (d["valor_venta"] - d["COSTO"]).round(2)
        resultados["utilidad_total_por_producto"] = d.groupby("producto")["UTILIDAD"].sum().sort_values(ascending=False)
        resultados["producto_mayor_utilidad"] = resultados["utilidad_total_por_producto"].idxmax()
        resultados["dataset_con_utilidad"] = d
    else:
        resultados["utilidad_total_por_producto"] = pd.Series(dtype=float)
        resultados["producto_mayor_utilidad"] = None
        resultados["dataset_con_utilidad"] = d

    return resultados

# -------------------- Extras groupby del enunciado --------------------
def extras_groupby(df: pd.DataFrame) -> Dict[str, pd.DataFrame]:
    """Genera las tablas extra solicitadas en el PDF con groupby()."""
    out: Dict[str, pd.DataFrame] = {}

    # Promedio de comisión por vendedor
    if "vendedor" in df.columns and "comision" in df.columns:
        out["promedio_comision_por_vendedor"] = (
            df.groupby("vendedor")["comision"].mean().sort_values(ascending=False).round(2).reset_index()
        )
    else:
        out["promedio_comision_por_vendedor"] = pd.DataFrame(columns=["vendedor", "comision"])

    # Ventas por estado y ciudad
    if "estado" in df.columns and "ciudad" in df.columns and "valor_venta" in df.columns:
        out["ventas_por_estado_y_ciudad"] = (
            df.groupby(["estado", "ciudad"])["valor_venta"].sum().reset_index()
            .sort_values(["estado", "valor_venta"], ascending=[True, False])
        )
    else:
        out["ventas_por_estado_y_ciudad"] = pd.DataFrame(columns=["estado", "ciudad", "valor_venta"])

    # Valor total por categoría
    if "categoria" in df.columns and "valor_venta" in df.columns:
        out["valor_total_por_categoria"] = (
            df.groupby("categoria")["valor_venta"].sum().sort_values(ascending=False).reset_index()
        )
    else:
        out["valor_total_por_categoria"] = pd.DataFrame(columns=["categoria", "valor_venta"])

    # Cerradas por vendedor y ciudad
    if "estado" in df.columns and "vendedor" in df.columns and "ciudad" in df.columns:
        cerr = df[df["estado"] == "Cerrado"]
        out["cerradas_por_vendedor_y_ciudad"] = (
            cerr.groupby(["vendedor", "ciudad"]).size().reset_index(name="cerradas")
            .sort_values("cerradas", ascending=False)
        )
    else:
        out["cerradas_por_vendedor_y_ciudad"] = pd.DataFrame(columns=["vendedor", "ciudad", "cerradas"])

    return out

# -------------------- Guardado --------------------
def _to_csv(obj: Any, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if isinstance(obj, (pd.Series, pd.DataFrame)):
        # Series: conservar índice con nombre claro
        if isinstance(obj, pd.Series):
            df = obj.reset_index()
            if df.columns[0] == "index":
                df = df.rename(columns={"index": "clave"})
            df.to_csv(path, index=False)
        else:
            obj.to_csv(path, index=False)
    else:
        Path(path).write_text(str(obj), encoding="utf-8")

def guardar_resultados(resultados: Dict[str, Any], extras: Dict[str, pd.DataFrame], out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)

    # principales
    if "ventas_por_estado" in resultados:
        _to_csv(resultados["ventas_por_estado"], out_dir / "ventas_por_estado.csv")
    if "ventas_por_ciudad_sum" in resultados:
        s = resultados["ventas_por_ciudad_sum"]
        df = s.reset_index().rename(columns={"index": "CIUDAD", "valor_venta": "VALOR_VENTA"})
        if "ciudad" in df.columns:
            df = df.rename(columns={"ciudad": "CIUDAD"})
        _to_csv(df, out_dir / "ventas_por_ciudad_sum.csv")

    if "top5_productos_por_registros" in resultados:
        s = resultados["top5_productos_por_registros"].reset_index()
        s.columns = ["PRODUCTO", "CUENTA"]
        _to_csv(s, out_dir / "top5_productos_por_registros.csv")

    if "media_valor_venta_por_mes" in resultados:
        _to_csv(resultados["media_valor_venta_por_mes"].reset_index().rename(columns={"mes": "Mes", "valor_venta": "MEDIA"}), out_dir / "media_valor_venta_por_mes.csv")

    if "ventas_por_trimestre" in resultados:
        tri = pd.Series(resultados["ventas_por_trimestre"], name="CUENTA").reset_index().rename(columns={"index": "TRIMESTRE"})
        _to_csv(tri, out_dir / "ventas_por_trimestre.csv")

    if "productos_mas_de_3_ciudades" in resultados:
        s = resultados["productos_mas_de_3_ciudades"].reset_index()
        s.columns = ["PRODUCTO", "N_CIUDADES"]
        _to_csv(s, out_dir / "productos_mas_de_3_ciudades.csv")

    if "utilidad_total_por_producto" in resultados:
        _to_csv(resultados["utilidad_total_por_producto"].reset_index().rename(columns={"index": "PRODUCTO", "UTILIDAD": "UTILIDAD_TOTAL"}), out_dir / "utilidad_total_por_producto.csv")

    # chequeos de calidad
    neg = pd.DataFrame({
        "ventas_valor_negativo": [resultados.get("ventas_valor_negativo", 0)],
        "comision_negativa": [resultados.get("comision_negativa", 0)],
        "nulos_valor": [resultados.get("nulos_valor", 0)],
        "nulos_comision": [resultados.get("nulos_comision", 0)],
    })
    _to_csv(neg, out_dir / "ventas_valor_o_comision_no_positiva.csv")

    # extras
    for name, df in extras.items():
        _to_csv(df, out_dir / f"{name}.csv")

    # derivados
    if "ventas_mensuales_por_ciudad" in resultados:
        _to_csv(resultados["ventas_mensuales_por_ciudad"], out_dir / "ventas_mensuales_por_ciudad.csv")

    # dataset con columnas agregadas
    if "dataset_con_utilidad" in resultados:
        _to_csv(resultados["dataset_con_utilidad"], out_dir / "dataset_con_utilidad.csv")

    # resumen de texto
    resumen_lines = []
    resumen_lines.append(f"Total de registros: {resultados.get('total_registros', 0)}")
    resumen_lines.append("Ventas por estado:")
    if isinstance(resultados.get("ventas_por_estado"), pd.Series):
        resumen_lines.append(resultados["ventas_por_estado"].to_string())
    resumen_lines.append(f"Valor total de ventas: {fmt_cop(resultados.get('valor_total_ventas', 0.0))}")
    resumen_lines.append(f"Promedio comisión (cerradas): {fmt_cop(resultados.get('prom_comision_cerradas', 0.0))}")
    resumen_lines.append(f"Ciudad con más ventas cerradas: {resultados.get('ciudad_mas_cerradas')}")
    resumen_lines.append("Top 5 productos por número de registros:")
    if isinstance(resultados.get("top5_productos_por_registros"), pd.Series):
        resumen_lines.append(resultados["top5_productos_por_registros"].to_string())
    resumen_lines.append(f"Productos únicos: {resultados.get('productos_unicos', 0)}")
    resumen_lines.append(f"Vendedor con más ventas cerradas: {resultados.get('vendedor_mas_cerradas')}")
    if isinstance(resultados.get("venta_max"), (pd.Series, pd.DataFrame)):
        resumen_lines.append("Venta de mayor valor:")
        resumen_lines.append(str(resultados["venta_max"]))
    resumen_lines.append(f"Duplicados totales: {resultados.get('duplicados_total', 0)}")
    resumen_lines.append(f"Filas eliminadas por nulos (cliente/producto/valor_venta): {resultados.get('filas_eliminadas_por_nulos', 0)}")
    resumen_lines.append(f"Producto con mayor 'UTILIDAD' total (según enunciado): {resultados.get('producto_mayor_utilidad')}")

    (out_dir / "resumen.txt").write_text("\n".join(resumen_lines), encoding="utf-8")

# -------------------- Informe Word (opcional) --------------------
def generar_informe_word(df: pd.DataFrame, resultados: Dict[str, Any], extras: Dict[str, pd.DataFrame], out_dir: Path) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        raise RuntimeError("python-docx no está instalado. Ejecuta: pip install python-docx") from e

    doc = Document()
    title = doc.add_heading("Informe Analítico de Ventas (Pandas)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M"))

    ventas_por_estado = resultados.get("ventas_por_estado", pd.Series(dtype=int))

    venta_max = resultados.get("venta_max")
    doc.add_paragraph(
        f"Se analizaron {len(df):,} registros de ventas, con un valor total de {fmt_cop(resultados.get('valor_total_ventas', 0.0))}. "
        f"El estado 'Cerrado' representa {int(ventas_por_estado.get('Cerrado', 0)):,} transacciones; la comisión promedio para ventas cerradas es "
        f"{fmt_cop(resultados.get('prom_comision_cerradas', 0.0))}. La ciudad con mayor número de ventas cerradas fue {resultados.get('ciudad_mas_cerradas')}. "
        f"El vendedor con mayor número de ventas cerradas fue {resultados.get('vendedor_mas_cerradas')}."
    )

    # Ventas por estado (tabla)
    doc.add_heading("KPIs y Distribuciones", level=2)
    doc.add_paragraph("Ventas por estado:")
    t = doc.add_table(rows=1, cols=2)
    hdr = t.rows[0].cells
    hdr[0].text = "Estado"; hdr[1].text = "Cuenta"
    for k, v in ventas_por_estado.items():
        row = t.add_row().cells
        row[0].text = str(k); row[1].text = f"{v:,}"

    # Valor total por ciudad (Top 10)
    doc.add_paragraph("\nValor total de ventas por ciudad (Top 10):")
    val_city = resultados.get("ventas_por_ciudad_sum", pd.Series(dtype=float)).head(10)
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Ciudad"; t.rows[0].cells[1].text = "Valor total"
    for city, val in val_city.items():
        r = t.add_row().cells
        r[0].text = str(city); r[1].text = fmt_cop(val)

    # Top 5 productos
    doc.add_paragraph("\nTop 5 productos por número de registros:")
    top5 = resultados.get("top5_productos_por_registros", pd.Series(dtype=int))
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Producto"; t.rows[0].cells[1].text = "Registros"
    for prod, cnt in top5.items():
        r = t.add_row().cells
        r[0].text = str(prod); r[1].text = f"{int(cnt):,}"

    # Ventas por trimestre
    doc.add_paragraph("\nVentas por trimestre:")
    tri = resultados.get("ventas_por_trimestre", {})
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Trimestre"; t.rows[0].cells[1].text = "Cuenta"
    for tri_key, cnt in tri.items():
        r = t.add_row().cells
        r[0].text = str(tri_key); r[1].text = f"{cnt:,}"

    # Cobertura y portafolio
    doc.add_heading("Cobertura y Portafolio", level=2)
    doc.add_paragraph("Productos vendidos en más de 3 ciudades (Top 10):")
    prod_ci = resultados.get("productos_mas_de_3_ciudades", pd.Series(dtype=int))
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Producto"; t.rows[0].cells[1].text = "Ciudades distintas"
    for prod, nci in prod_ci.head(10).items():
        r = t.add_row().cells
        r[0].text = str(prod); r[1].text = str(int(nci))

    # Aclaración de "utilidad"
    doc.add_heading("'Utilidad' según enunciado", level=2)
    doc.add_paragraph(
        "Se creó la columna UTILIDAD = 95% del VALOR_VENTA (el enunciado la denomina 'utilidad' aunque operativamente equivale a un costo simulado). "
        "Se agregan también COSTO (95%) y UTILIDAD_5PCT (5%) para claridad."
    )

    # Extras groupby (resumen en tablas cortas)
    doc.add_heading("Extras groupby()", level=2)
    for name, df_extra in extras.items():
        doc.add_paragraph(f"\n{name} (primeras filas):")
        head = df_extra.head(10)
        t = doc.add_table(rows=head.shape[0] + 1, cols=head.shape[1])
        # header
        for j, col in enumerate(head.columns):
            t.cell(0, j).text = str(col)
        # rows
        for i in range(head.shape[0]):
            for j in range(head.shape[1]):
                t.cell(i + 1, j).text = str(head.iat[i, j])

    out_path = out_dir / "Informe_Analitico_Ventas_Pandas.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path

# -------------------- CLI --------------------
def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Ventas 360 — CLI (Pandas)")
    p.add_argument("--generate", action="store_true", help="Genera dataset sintético")
    p.add_argument("--analyze", action="store_true", help="Ejecuta análisis y exporta resultados")
    p.add_argument("--n", type=int, default=5000, help="Número de filas a generar (si --generate)")
    p.add_argument("--csv", type=Path, default=Path("data/ventas_5k.csv"), help="Ruta del CSV (entrada o salida de --generate)")
    p.add_argument("--out", type=Path, default=Path("out"), help="Directorio de salida")
    p.add_argument("--word", action="store_true", help="Generar informe Word (.docx)")
    p.add_argument("--seed", type=int, default=RANDOM_SEED, help="Seed aleatoria")
    return p

def main() -> None:
    args = build_parser().parse_args()
    csv_path: Path = args.csv
    out_dir: Path = args.out

    if args.generate:
        print(f"Generando dataset en {csv_path} ...")
        generar_dataset(args.n, csv_path, seed=args.seed)
        print("✅ CSV generado.")

    if args.analyze:
        print(f"Cargando CSV desde {csv_path} ...")
        df = cargar_csv(csv_path)
        print("Analizando con pandas...")
        resultados = responder_preguntas(df)
        extras = extras_groupby(df)
        print("Guardando resultados en", out_dir.resolve())
        guardar_resultados(resultados, extras, out_dir)

        if args.word:
            print("Generando Informe Administrativo (Word)...")
            out_docx = generar_informe_word(df, resultados, extras, out_dir)
            print("📄 Informe:", out_docx.resolve())

        print("✅ Análisis completado.")

if __name__ == "__main__":
    main()